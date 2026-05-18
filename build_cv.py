"""
Builds every CV artefact from the single source of truth, MASTER-CV.md:
  - Deepa_Palaniappan_CV_2026.docx   (curated 3-page Word CV; PDF made by build.ps1)
  - docs/resume.json                 (JSON Resume standard, full record)
  - docs/index.html                  (accessible web CV, full record)

Visibility markers in MASTER-CV.md:
  '- ' both    '+ ' web only    '~ ' PDF only    '> ' note (PDF only)
A '### + ...' / '### ~ ...' role applies that visibility to its bullets too.

You never edit this script for content. Edit MASTER-CV.md, then run build.ps1.
"""
import json
import re
import html
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE

HERE = Path(__file__).parent
SRC = HERE / 'MASTER-CV.md'

NAVY = RGBColor(0x1F, 0x29, 0x37)
ACCENT = RGBColor(0x1E, 0x40, 0xAF)
MUTED = RGBColor(0x4B, 0x55, 0x63)

PDF_ALLOW = {'both', 'pdf'}
WEB_ALLOW = {'both', 'web'}


# ---------------------------------------------------------------- parsing ---
def _marker(s):
    """Return (visibility, rest) for a content line, or (None, None)."""
    if s.startswith('- '):
        return 'both', s[2:].strip()
    if s.startswith('+ '):
        return 'web', s[2:].strip()
    if s.startswith('~ '):
        return 'pdf', s[2:].strip()
    return None, None


def parse(text):
    meta = {}
    sections = []
    cur = None
    in_cv = False
    joined = re.sub(r'<!--.*?-->', '', text, flags=re.S)
    for raw in joined.splitlines():
        s = raw.strip()
        if s == '# INBOX':
            in_cv = False
            continue
        if s == '# CV':
            in_cv = True
            continue
        if not in_cv:
            continue
        if not s:
            if cur and cur['kind'] == 'profile' and cur['items'] \
                    and cur['items'][-1] != '':
                cur['items'].append('')
            continue
        m = re.match(r'@(\w+):\s*(.*)', s)
        if m:
            meta[m.group(1)] = m.group(2).strip()
            continue
        if s.startswith('## '):
            title = s[3:].strip()
            up = title.upper()
            kind = ('profile' if up == 'PROFILE'
                    else 'skills' if 'AREAS OF EXPERTISE' in up
                    else 'experience' if 'EXPERIENCE' in up
                    else 'list')
            cur = {'title': title, 'kind': kind, 'items': []}
            sections.append(cur)
            continue
        if cur is None:
            continue
        if s.startswith('### '):
            body = s[4:].strip()
            vis = 'both'
            if body.startswith('+ '):
                vis, body = 'web', body[2:].strip()
            elif body.startswith('~ '):
                vis, body = 'pdf', body[2:].strip()
            bits = [x.strip() for x in body.split(' @@ ')]
            cur['items'].append({
                'role': bits[0] if bits else '',
                'org': bits[1] if len(bits) > 1 else '',
                'dates': bits[2] if len(bits) > 2 else '',
                'vis': vis, 'bullets': []})
            continue
        if s.startswith('> '):
            cur['items'].append({'note': s[2:].strip(), 'vis': 'pdf'})
            continue
        vis, body = _marker(s)
        if vis is None:
            if cur['kind'] == 'profile':
                cur['items'].append(s)
            continue
        if cur['kind'] == 'skills':
            indent = len(raw) - len(raw.lstrip(' '))
            if indent >= 2 and cur['items']:
                cur['items'][-1]['details'].append(body)
            else:
                cur['items'].append({'term': body, 'vis': vis,
                                     'details': []})
            continue
        if cur['kind'] == 'experience' and cur['items'] \
                and isinstance(cur['items'][-1], dict) \
                and 'bullets' in cur['items'][-1]:
            role = cur['items'][-1]
            bvis = vis if s[0] in '+~' else role['vis']
            role['bullets'].append((body, bvis))
        else:
            cur['items'].append((body, vis))
    for sec in sections:
        if sec['kind'] == 'profile':
            paras, buf = [], []
            for it in sec['items']:
                if it == '':
                    if buf:
                        paras.append(' '.join(buf))
                        buf = []
                else:
                    buf.append(it)
            if buf:
                paras.append(' '.join(buf))
            sec['items'] = paras
    return meta, sections


def bold_spans(text):
    out, bold = [], False
    for part in text.split('**'):
        if part:
            out.append((part, bold))
        bold = not bold
    return out or [(text, False)]


LINK_RE = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')


def inline_tokens(text):
    """Split text into (text, is_bold, href_or_None) runs.

    Supports **bold** and [label](url) markdown so one source file can
    produce links in Word, PDF and HTML.
    """
    tokens, pos = [], 0
    for m in LINK_RE.finditer(text):
        if m.start() > pos:
            tokens += [(t, b, None) for t, b in bold_spans(text[pos:m.start()])]
        tokens.append((m.group(1), False, m.group(2)))
        pos = m.end()
    if pos < len(text):
        tokens += [(t, b, None) for t, b in bold_spans(text[pos:])]
    return tokens or [(text, False, None)]


def strip_md(text):
    """Plain text for JSON: drop link syntax and bold markers."""
    return LINK_RE.sub(r'\1', text).replace('**', '')


def visible(sections, allow):
    """Return sections filtered to the items allowed for this output."""
    res = []
    for sx in sections:
        if sx['kind'] == 'profile':
            res.append(sx)
            continue
        items = []
        for it in sx['items']:
            if isinstance(it, tuple):
                if it[1] in allow:
                    items.append(it)
            elif 'note' in it:
                if it['vis'] in allow:
                    items.append(it)
            elif 'term' in it:
                if it['vis'] in allow:
                    items.append(it)
            else:  # role
                if it['vis'] in allow:
                    b = [x for x in it['bullets'] if x[1] in allow]
                    items.append({**it, 'bullets': b})
        if items:
            res.append({**sx, 'items': items})
    return res


# ------------------------------------------------------------------ docx ---
def build_docx(meta, sections, out_path):
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Calibri'
    st.font.size = Pt(10.5)
    st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    st.paragraph_format.space_before = Pt(0)
    st.paragraph_format.space_after = Pt(0)
    st.paragraph_format.line_spacing = 1.1

    sec = doc.sections[0]
    sec.page_height = Inches(11.69)
    sec.page_width = Inches(8.27)
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.65)
    sec.right_margin = Inches(0.65)

    def para(sa=2, sb=0, align=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(sa)
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.line_spacing = 1.1
        if align is not None:
            p.alignment = align
        return p

    def run(p, t, size=10.5, bold=False, italic=False, color=None):
        r = p.add_run(t)
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        if color is not None:
            r.font.color.rgb = color
        return r

    def heading(t):
        p = para(sa=3, sb=9)
        run(p, t.upper(), size=11, bold=True, color=NAVY)
        pPr = p._p.get_or_add_pPr()
        bdr = OxmlElement('w:pBdr')
        b = OxmlElement('w:bottom')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '6')
        b.set(qn('w:space'), '1')
        b.set(qn('w:color'), '1F2937')
        bdr.append(b)
        pPr.append(bdr)

    def add_link(p, url, text, size=10.5):
        r_id = p.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK,
                                is_external=True)
        link = OxmlElement('w:hyperlink')
        link.set(qn('r:id'), r_id)
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        col = OxmlElement('w:color')
        col.set(qn('w:val'), '1E40AF')
        rPr.append(col)
        und = OxmlElement('w:u')
        und.set(qn('w:val'), 'single')
        rPr.append(und)
        szc = OxmlElement('w:sz')
        szc.set(qn('w:val'), str(int(size * 2)))
        rPr.append(szc)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = text
        r.append(t)
        link.append(r)
        p._p.append(link)

    def bullet(text, sa=2):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(sa)
        pf.space_before = Pt(0)
        pf.line_spacing = 1.1
        pf.left_indent = Inches(0.16)
        pf.first_line_indent = Inches(-0.16)
        rb = p.add_run('•  ')
        rb.font.size = Pt(10.5)
        for t, b, href in inline_tokens(text):
            if href:
                add_link(p, href, t)
            else:
                r = p.add_run(t)
                r.font.size = Pt(10.5)
                r.bold = b

    p = para(sa=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(p, meta.get('name', ''), size=20, bold=True, color=NAVY)
    p = para(sa=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(p, '   |   '.join(x.strip() for x in meta.get('headline', '').split('|')),
        size=10.5, bold=True, color=ACCENT)
    p = para(sa=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(p, meta.get('contact', ''), size=9, color=MUTED)

    for sx in sections:
        heading(sx['title'])
        if sx['kind'] == 'profile':
            for ptxt in sx['items']:
                run(para(sa=2), ptxt, size=10.5)
        elif sx['kind'] == 'skills':
            items = [it['term'] for it in sx['items']]
            for i in range(0, len(items), 2):
                pair = items[i:i + 2]
                pp = doc.add_paragraph()
                pf = pp.paragraph_format
                pf.space_after = Pt(2)
                pf.space_before = Pt(0)
                pf.line_spacing = 1.1
                pf.left_indent = Inches(0.16)
                pf.first_line_indent = Inches(-0.16)
                r = pp.add_run('•  ' + pair[0])
                r.font.size = Pt(10.5)
                if len(pair) > 1:
                    r2 = pp.add_run('      •  ' + pair[1])
                    r2.font.size = Pt(10.5)
        elif sx['kind'] == 'experience':
            for it in sx['items']:
                run(para(sa=1, sb=5), it['role'], size=10.5, bold=True,
                    color=NAVY)
                meta_bits = [x for x in (it.get('org', ''),
                                         it.get('dates', '')) if x]
                if meta_bits:
                    run(para(sa=2), '  ·  '.join(meta_bits), size=9,
                        italic=True, color=MUTED)
                for b, _ in it['bullets']:
                    bullet(b)
        else:
            for it in sx['items']:
                if isinstance(it, dict) and 'note' in it:
                    run(para(sa=2, sb=2), it['note'], size=9.5, italic=True,
                        color=MUTED)
                else:
                    bullet(it[0])

    doc.save(str(out_path))


# -------------------------------------------------------------- resume.json ---
def build_resume_json(meta, sections, out_path):
    parts = [c.strip() for c in re.split(r'·|\|', meta.get('contact', ''))]
    email = next((c for c in parts if '@' in c and ' ' not in c), '')
    phone = next((c for c in parts if re.search(r'\d{5,}', c)), '')
    site = next((c for c in parts if '.' in c and '@' not in c
                 and 'bit.ly' not in c), '')
    loc = next((c for c in parts if ',' in c), '')
    basics = {'name': meta.get('name', ''),
              'label': meta.get('headline', '').split('|')[0].strip(),
              'email': email, 'phone': phone,
              'url': ('https://' + site) if site
              and not site.startswith('http') else site,
              'summary': '', 'location': {'address': loc}, 'profiles': []}
    work, education, publications, awards, skills = [], [], [], [], []
    for sx in sections:
        up = sx['title'].upper()
        if sx['kind'] == 'profile':
            basics['summary'] = ' '.join(sx['items'])
        elif sx['kind'] == 'experience':
            for it in sx['items']:
                work.append({'name': strip_md(it.get('org', '')),
                             'position': strip_md(it['role']),
                             'description': it.get('dates', ''),
                             'highlights': [strip_md(b)
                                            for b, _ in it['bullets']]})
        elif 'EDUCATION' in up:
            education = [{'institution': strip_md(it[0])}
                         for it in sx['items'] if isinstance(it, tuple)]
        elif 'PUBLICATION' in up:
            publications = [{'name': strip_md(it[0])}
                            for it in sx['items'] if isinstance(it, tuple)]
        elif 'FELLOWSHIP' in up or 'CERTIF' in up or 'COMMITTEE' in up:
            awards = [{'title': strip_md(it[0])} for it in sx['items']
                      if isinstance(it, tuple)]
        elif sx['kind'] == 'skills':
            skills = [{'name': strip_md(it['term'])}
                      for it in sx['items']]
    data = {'$schema': 'https://raw.githubusercontent.com/jsonresume/'
            'resume-schema/v1.0.0/schema.json',
            'basics': basics, 'work': work, 'education': education,
            'publications': publications, 'awards': awards, 'skills': skills}
    out_path.write_text(json.dumps(data, indent=2, ensure_ascii=False),
                         encoding='utf-8')


# ----------------------------------------------------------------- index.html ---
def linkify_contact(contact):
    out = []
    for tok in [c.strip() for c in contact.split('·') if c.strip()]:
        if '@' in tok and ' ' not in tok:
            out.append(f'<a href="mailto:{html.escape(tok)}">'
                       f'{html.escape(tok)}</a>')
        elif re.fullmatch(r'\+?[\d ]{7,}', tok):
            out.append(f'<a href="tel:{tok.replace(" ", "")}">'
                       f'{html.escape(tok)}</a>')
        elif '.' in tok and ' ' not in tok:
            url = tok if tok.startswith('http') else 'https://' + tok
            out.append(f'<a href="{html.escape(url)}">{html.escape(tok)}</a>')
        else:
            out.append(f'<span>{html.escape(tok)}</span>')
    return ' <span aria-hidden="true">·</span> '.join(out)


def inline_html(text):
    out = []
    for t, b, href in inline_tokens(text):
        esc = html.escape(t)
        if href:
            out.append(f'<a href="{html.escape(href)}">{esc}</a>')
        elif b:
            out.append(f'<strong>{esc}</strong>')
        else:
            out.append(esc)
    return ''.join(out)


def build_html(meta, sections, out_path):
    name = html.escape(meta.get('name', ''))
    headline = html.escape(meta.get('headline', ''))
    body = []
    toc = []
    for sx in sections:
        sid = re.sub(r'[^a-z0-9]+', '-', sx['title'].lower()).strip('-')
        toc.append((sx['title'], sid))
        body.append(f'<section aria-labelledby="{sid}">')
        body.append(f'<h2 id="{sid}">{html.escape(sx["title"])}</h2>')
        if sx['kind'] == 'profile':
            body += [f'<p>{html.escape(p)}</p>' for p in sx['items']]
        elif sx['kind'] == 'skills':
            body.append('<ul class="exp" role="list">')
            for it in sx['items']:
                term = html.escape(it['term'])
                if it.get('details'):
                    body.append('<li><details><summary>'
                                f'{term}</summary><ul>')
                    body += [f'<li>{inline_html(d)}</li>'
                             for d in it['details']]
                    body.append('</ul></details></li>')
                else:
                    body.append(f'<li>{term}</li>')
            body.append('</ul>')
        elif sx['kind'] == 'experience':
            for it in sx['items']:
                body.append('<article class="role">')
                body.append(f'<h3>{inline_html(it["role"])}</h3>')
                meta_bits = [html.escape(x) for x in
                             (it.get('org', ''), it.get('dates', '')) if x]
                if meta_bits:
                    body.append('<p class="dates">'
                                + ' &middot; '.join(meta_bits) + '</p>')
                if it['bullets']:
                    body.append('<ul>')
                    body += [f'<li>{inline_html(b)}</li>'
                             for b, _ in it['bullets']]
                    body.append('</ul>')
                body.append('</article>')
        else:
            note = None
            body.append('<ul>')
            for it in sx['items']:
                if isinstance(it, dict) and 'note' in it:
                    note = it['note']
                else:
                    body.append(f'<li>{inline_html(it[0])}</li>')
            body.append('</ul>')
            if note:
                body.append(f'<p class="note">{html.escape(note)}</p>')
        body.append('</section>')

    toc_links = '\n'.join(
        f'  <li><a href="#{sid}">{html.escape(t)}</a></li>'
        for t, sid in toc)
    nav_html = ('<nav class="toc" aria-labelledby="toc-h">\n'
                '<h2 id="toc-h">On this page</h2>\n<ul>\n'
                f'{toc_links}\n</ul>\n</nav>')

    doc = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{name} CV</title>
<meta name="description" content="{headline}">
<style>
  :root {{ --ink:#0b0c0c; --muted:#4b5563; --accent:#1e40af; --rule:#d1d5db; }}
  * {{ box-sizing:border-box; }}
  body {{ margin:0; font:1.05rem/1.6 -apple-system,BlinkMacSystemFont,
    "Segoe UI",Roboto,Helvetica,Arial,sans-serif; color:var(--ink);
    background:#fff; }}
  .skip {{ position:absolute; left:-999px; top:0; background:#fde68a;
    color:#78350f; padding:.6rem 1rem; }}
  .skip:focus {{ left:0; }}
  main {{ max-width:48rem; margin:0 auto; padding:2.2rem 1.4rem 4rem; }}
  header.cv {{ text-align:center; border-bottom:3px solid var(--ink);
    padding-bottom:1rem; margin-bottom:1.4rem; }}
  h1 {{ font-size:2.2rem; margin:.2rem 0; }}
  .headline {{ color:var(--accent); font-weight:700; }}
  .contact {{ color:var(--muted); font-size:.95rem; margin-top:.5rem; }}
  h2 {{ font-size:1.3rem; border-bottom:1px solid var(--rule);
    padding-bottom:.25rem; margin:2.4rem 0 .8rem;
    scroll-margin-top:.5rem; }}
  h3 {{ font-size:1rem; margin:1rem 0 .15rem; }}
  .dates {{ color:var(--muted); font-size:.9rem; font-style:italic;
    margin:.1rem 0 .4rem; }}
  ul {{ margin:.3rem 0 .6rem; padding-left:1.2rem; }}
  li {{ margin:.25rem 0; }}
  ul.exp {{ list-style:none; padding:0; }}
  ul.exp > li {{ margin:.35rem 0; }}
  details {{ border-bottom:1px solid var(--rule); padding:.1rem 0; }}
  summary {{ cursor:pointer; font-weight:600; padding:.5rem 0;
    color:var(--accent); }}
  summary:focus-visible {{ outline:3px solid #78350f;
    background:#fde68a; color:#0b0c0c; }}
  details[open] summary {{ margin-bottom:.2rem; }}
  details ul {{ margin:.2rem 0 .6rem; }}
  nav.toc {{ border:1px solid var(--rule); border-radius:.4rem;
    padding:.6rem 1rem; margin:0 0 1.8rem; background:#f8f9fb; }}
  nav.toc h2 {{ border:0; margin:.2rem 0 .5rem; font-size:1rem;
    padding:0; }}
  nav.toc ul {{ list-style:none; padding:0; margin:0; display:flex;
    flex-wrap:wrap; gap:.35rem 1.1rem; }}
  nav.toc li {{ margin:0; }}
  .note {{ color:var(--muted); font-style:italic; font-size:.92rem; }}
  a {{ color:var(--accent); }}
  a:focus-visible {{ outline:3px solid #78350f; background:#fde68a;
    color:#0b0c0c; }}
  .actions {{ text-align:center; margin-bottom:1.5rem; }}
  .actions a {{ display:inline-block; border:1px solid var(--accent);
    padding:.4rem .9rem; border-radius:.3rem; text-decoration:none;
    font-size:.95rem; }}
  @media (prefers-reduced-motion: reduce) {{ * {{ animation:none !important;
    transition:none !important; }} }}
  @media print {{ .skip,.actions,nav.toc {{ display:none; }}
    details > * {{ display:block; }}
    summary {{ font-weight:700; color:var(--ink); }}
    a {{ color:var(--ink); text-decoration:none; }}
    body {{ font-size:10.5pt; }} }}
</style>
</head>
<body>
<a class="skip" href="#main">Skip to content</a>
<main id="main">
<header class="cv">
  <h1>{name}</h1>
  <p class="headline">{headline}</p>
  <p class="contact">{linkify_contact(meta.get('contact', ''))}</p>
</header>
{nav_html}
<p class="actions">
  <a href="Deepa_Palaniappan_CV_2026.pdf">Download CV (PDF, 3 pages)</a>
</p>
{chr(10).join(body)}
</main>
</body>
</html>
"""
    out_path.write_text(doc, encoding='utf-8')


def main():
    meta, sections = parse(SRC.read_text(encoding='utf-8'))
    docs = HERE / 'docs'
    docs.mkdir(exist_ok=True)
    build_docx(meta, visible(sections, PDF_ALLOW),
               HERE / 'Deepa_Palaniappan_CV_2026.docx')
    web = visible(sections, WEB_ALLOW)
    build_resume_json(meta, web, docs / 'resume.json')
    build_html(meta, web, docs / 'index.html')
    n_pdf = sum(len(s['items']) for s in visible(sections, PDF_ALLOW)
                if s['kind'] != 'profile')
    n_web = sum(len(s['items']) for s in web if s['kind'] != 'profile')
    print('Built: Deepa_Palaniappan_CV_2026.docx (root), '
          'docs/index.html, docs/resume.json')
    print(f'Curated PDF items: {n_pdf}   Full web items: {n_web}')


if __name__ == '__main__':
    main()
